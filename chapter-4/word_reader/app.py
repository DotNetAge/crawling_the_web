# coding:utf-8
import docx

doc = docx.Document('demo.docx')
print "文档中共发现%s个段落" % len(doc.paragraphs)

if len(doc.paragraphs):
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    print '\n'.join(fullText)
